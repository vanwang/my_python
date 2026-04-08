import os
import re
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from pathlib import Path

def count_chinese_chars(file_path):
    """
    统计Mac下docx文档的纯中文字数（仅匹配[\u4e00-\u9fff]，排除标点/字母/数字/空格）
    提取段落+表格文字，避免漏统计，返回整数字数
    """
    all_text = ""
    file_path = Path(file_path)
    try:
        # 仅处理docx格式（Mac原生不支持doc二进制格式）
        doc = Document(file_path)
        # 提取所有段落文字
        for para in doc.paragraphs:
            all_text += para.text.strip()
        # 提取表格中的所有文字（关键：避免表格内容漏统计）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text.strip()
        # 正则精准匹配所有中文字符（Unicode中文范围：\u4e00-\u9fff）
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', all_text)
        return len(chinese_chars)  # 返回纯中文字数（整数）

    except Exception as e:
        print(f"❌ 处理文档失败：{file_path.name}，错误：{str(e)}")
        return 0

def scan_word_files(dir_path):
    """
    遍历Mac指定目录，获取所有.docx文件（排除隐藏文件/文件夹）
    返回格式：[(文档标题, 绝对路径), ...]，标题为去掉后缀的文件名
    """
    word_files = []
    # 转换为绝对路径，适配Mac路径格式（/Users/xxx/文档 这种）
    dir_path = Path(dir_path).absolute()
    # 校验目录是否存在
    if not dir_path.is_dir():
        print(f"❌ 错误：指定目录不存在 → {dir_path}")
        return word_files
    # 遍历目录下所有文件（不递归子目录，如需递归请改为rglob("*.docx")）
    for file in dir_path.glob("*.docx"):
        # 排除Mac隐藏文件（以.开头的文件，如.DS_Store）
        if file.is_file() and not file.name.startswith("."):
            title = file.name.rsplit(".", 1)[0]  # 去掉.docx后缀作为标题
            word_files.append((title, str(file)))
    return word_files

def generate_excel(data, save_path):
    """
    生成Excel文件，包含「编号、标题、字数」三列
    美化表格：表头加粗居中、内容居中、设置合理列宽，Mac/WindowsExcel均可打开
    """
    # 创建工作簿，激活工作表
    wb = Workbook()
    ws = wb.active
    ws.title = "Word中文字数统计"  # Excel工作表名称

    # 设置表头
    headers = ["编号", "标题", "字数"]
    ws.append(headers)

    # 写入统计数据：编号自增，标题+字数对应统计结果
    for idx, (title, char_count) in enumerate(data, 1):
        ws.append([idx, title, char_count])

    # 表格美化（直接可用，无需二次调整）
    # 1. 表头：加粗、12号字、水平+垂直居中
    for cell in ws[1]:
        cell.font = Font(bold=True, size=12)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    # 2. 内容：所有单元格居中
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=3):
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center")
    # 3. 设置列宽（适配长标题，编号8、标题45、字数10）
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 45
    ws.column_dimensions["C"].width = 10

    # 保存Excel，处理保存异常（如路径无权限、文件被占用）
    try:
        # 确保保存目录存在（如指定/Users/xxx/统计结果/，自动创建该文件夹）
        save_dir = Path(save_path).parent
        if not save_dir.exists():
            save_dir.mkdir(parents=True, exist_ok=True)
        wb.save(save_path)
        print(f"\n✅ Excel统计文件生成成功 → {save_path}")
    except PermissionError:
        print(f"❌ 生成Excel失败：无权限写入路径 {save_path}，请检查路径权限")
    except Exception as e:
        print(f"❌ 生成Excel失败：{str(e)}")

def main():
    """主函数：Mac环境入口，仅需修改下方2个配置项"""
    # -------------------------- Mac路径配置（必改！）--------------------------
    TARGET_DIR = r"/Users/Van/Downloads/翻译"  # 要统计的docx所在目录（桌面/Word文档 示例）
    EXCEL_SAVE_PATH = r"/Users/Van/Downloads/翻译/Word字数统计结果.xlsx"  # Excel保存路径（桌面示例）
    # --------------------------------------------------------------------------

    # 1. 扫描目录下的docx文件
    print(f"🔍 开始扫描目录：{TARGET_DIR}")
    word_files = scan_word_files(TARGET_DIR)
    if not word_files:
        print("📭 未找到任何.docx文件，请检查目录是否正确！")
        return
    print(f"✅ 共找到 {len(word_files)} 个docx文档，开始统计中文字数...\n")

    # 2. 逐个统计纯中文字数
    stat_data = []
    for title, file_path in word_files:
        char_count = count_chinese_chars(file_path)
        stat_data.append((title, char_count))
        print(f"📄 {title} → 纯中文字数：{char_count} 字")

    # 3. 生成Excel统计文件
    generate_excel(stat_data, EXCEL_SAVE_PATH)

if __name__ == "__main__":
    # Mac终端打印中文防乱码
    os.environ["PYTHONIOENCODING"] = "utf-8"
    main()